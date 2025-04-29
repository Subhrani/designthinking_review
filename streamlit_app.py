import streamlit as st
from recode import (
    extract_sections,
    load_pdf_content,
    get_section_type,
    generate_mcq,
    generate_questions,
    generate_exam_docx,
)
from docx import Document as DocxDocument
import tempfile
import os

st.set_page_config(page_title="Quiz & Question Paper Generator", layout="wide")
st.title("Quiz & Question Paper Generator")

# --- CACHED PDF LOADER ---
@st.cache_data(show_spinner="Loading PDF... This may take a moment.")
def load_pdf_once(pdf_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(pdf_bytes)
        tmp_pdf_path = tmp_file.name
    pages = load_pdf_content(tmp_pdf_path)
    return pages, tmp_pdf_path

# 1. Upload PDF
st.header("Step 1: Upload PDF")
uploaded_file = st.file_uploader("Upload your PDF file", type=["pdf"])
pages = None
tmp_pdf_path = None
if uploaded_file:
    pages, tmp_pdf_path = load_pdf_once(uploaded_file.read())
    st.success("PDF uploaded and loaded!")

# 2. Select Section Split Method
if pages:
    st.header("Step 2: Select Section Split Method")
    section_type = st.selectbox(
        "How would you like to split the content?",
        ["chapter", "unit", "module", "section", "topic", "custom", "any"]
    )
    custom_keyword = None
    if section_type == "custom":
        custom_keyword = st.text_input("Enter custom keyword for splitting:")

    # 3. Extract Sections
    if st.button("Extract Sections"):
        full_text = "\n\n".join([page.page_content for page in pages])
        documents = extract_sections(full_text, section_type, custom_keyword)
        st.session_state['documents'] = documents
        st.success(f"Extracted {len(documents)} sections!")

# 4. Preview & Select Sections
if 'documents' in st.session_state:
    st.header("Step 3: Preview & Select Sections")
    documents = st.session_state['documents']
    section_titles = []
    for i, doc in enumerate(documents):
        first_line = doc.page_content.split('\n')[0][:100]
        section_titles.append(f"{i+1}. {doc.metadata.get('section', '')}: {first_line}...")
    selected_indices = st.multiselect("Select sections to generate questions from:", options=list(range(len(documents))), format_func=lambda i: section_titles[i])
    selected_docs = [documents[i] for i in selected_indices]

    # 5. Choose Assessment Type
    st.header("Step 4: Choose Assessment Type")
    assessment_type = st.selectbox("Assessment Type", ["quiz", "question paper"])

    # 6. Course Details
    st.header("Step 5: Enter Course Details")
    course = st.text_input("Course Name")
    code = st.text_input("Course Code")
    exam = st.text_input("Exam Title")

    # 7. Generate Quiz or Question Paper
    if assessment_type == "quiz":
        num_questions = st.number_input("Number of MCQs per section", min_value=1, value=5)
        marks_per_question = st.number_input("Marks per question", min_value=1, value=1)
        if st.button("Generate Quiz"):
            docx = DocxDocument()
            docx.add_heading(f"{exam} Quiz", 0)
            docx.add_paragraph(f"Course Name: {course}")
            docx.add_paragraph(f"Course Code: {code}")
            docx.add_paragraph(f"Total Marks: {num_questions * marks_per_question}")
            docx.add_paragraph("")
            question_num = 1
            for doc in selected_docs:
                docx.add_heading(doc.metadata['section'], level=1)
                mcqs = generate_mcq(doc.page_content, num_questions)
                for mcq in mcqs:
                    lines = mcq.split('\n')
                    question = lines[0].lstrip('Q) ')
                    options = lines[1:5]
                    correct = lines[5].lstrip('CORRECT: ') if len(lines) > 5 else 'A'
                    q_par = docx.add_paragraph(f"{question_num}. {question}")
                    for opt in options:
                        docx.add_paragraph(opt.strip(), style='List Bullet')
                    docx.add_paragraph(f"[{marks_per_question} marks] (Recommended time: 60 seconds)")
                    docx.add_paragraph("")
                    question_num += 1
            # Download
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out_file:
                docx.save(out_file.name)
                st.success("Quiz generated!")
                with open(out_file.name, "rb") as f:
                    st.download_button("Download Quiz DOCX", f, file_name="Generated_Quiz.docx")
    else:
        st.header("Question Paper Pattern")
        st.info("Specify the number of questions and marks for each cognitive level.")
        pattern = {}
        levels = ["Remember", "Understand", "Apply"]
        for level in levels:
            count = st.number_input(f"Number of {level} Questions", min_value=0, value=2, key=f"count_{level}")
            marks = st.number_input(f"Marks per {level} Question", min_value=1, value=2, key=f"marks_{level}")
            if count > 0:
                pattern[level] = (count, marks)
        if st.button("Generate Question Paper"):
            questions = []
            total_marks = 0
            for doc in selected_docs:
                for level, (count, marks) in pattern.items():
                    qs = generate_questions(doc.page_content, level, count, marks)
                    questions.extend(qs)
                    total_marks += count * marks
            # Generate DOCX
            output_file = generate_exam_docx(course, code, exam, questions, pattern)
            st.success("Question paper generated!")
            with open(output_file, "rb") as f:
                st.download_button("Download Question Paper DOCX", f, file_name=output_file)

# Clean up temp PDF file (no longer needed since caching is used)
if tmp_pdf_path and os.path.exists(tmp_pdf_path):
    os.remove(tmp_pdf_path)


