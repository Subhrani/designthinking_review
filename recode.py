import os
import re
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_community.document_loaders import PyPDFLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_groq import ChatGroq
from langchain.chains import RetrievalQA
import time
from datetime import datetime
from typing import List, Dict
from langchain.chains import LLMChain
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate

# Load environment variables
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")

if not groq_api_key:
    raise RuntimeError("GROQ_API_KEY is not set. Please add it to your .env file.")

# Initialize Groq LLM once at the top
llm = ChatGroq(temperature=0.3, model_name="llama3-70b-8192")

def print_banner():
    """Print the application banner"""
    print("\n" + "="*60)
    print(" Question Paper Generator")
    print("="*60 + "\n")

def extract_sections(text, section_type="any", custom_keyword=None):
    """
    Extract sections (chapters/units/modules/sections/topics/custom) from text.
    section_type can be "chapter", "unit", "module", "section", "topic", "custom", or "any".
    If custom_keyword is provided, split by that keyword.
    """
    # Patterns for different section types
    patterns = {
        "chapter": r"(?:^|\n)(?:CHAPTER|Chapter)\s*[-–:]?\s*(\d+)[.\s]*([^\n]*)",
        "unit": r"(?:^|\n)(?:UNIT|Unit)\s*[-–:]?\s*(\d+)[.\s]*([^\n]*)",
        "module": r"(?:^|\n)(?:MODULE|Module)\s*[-–:]?\s*(\d+)[.\s]*([^\n]*)",
        "section": r"(?:^|\n)(?:SECTION|Section)\s*[-–:]?\s*(\d+)[.\s]*([^\n]*)",
        "topic": r"(?:^|\n)(?:TOPIC|Topic)\s*[-–:]?\s*(\d+)[.\s]*([^\n]*)",
    }
    
    if section_type == "custom" and custom_keyword:
        pattern = rf"(?:^|\n)({re.escape(custom_keyword)})[.\s]*([^\n]*)"
    elif section_type == "any":
        # Combine all patterns
        pattern = r"(?:^|\n)(?:CHAPTER|Chapter|UNIT|Unit|MODULE|Module|SECTION|Section|TOPIC|Topic)\s*[-–:]?\s*(\d+)[.\s]*([^\n]*)"
    else:
        pattern = patterns.get(section_type.lower())
        if pattern is None:
            raise ValueError(f"Unknown section_type: {section_type}")
    
    section_headers = list(re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE))
    section_positions = [(m.start(), m.group(1), m.group(2) if m.lastindex > 1 else "") for m in section_headers]
    
    def is_relevant(title, number):
        title_clean = title.strip()
        title_lower = title_clean.lower()
        # Ignore extremely high chapter numbers (likely noise)
        try:
            num = int(number)
            if num > 30:
                return False
        except Exception:
            return False
        # Ignore titles with file codes, dates, or page numbers
        ignore_patterns = [
            r"figure", r"table", r"answer key", r"solution", r"appendix", r"contents", r"index", r"glossary",
            r"\d{2,} [0-9a-zA-Z ]*ch\d+", # file code patterns
            r"\d{1,2}/\d{1,2}/\d{2,4}", # dates
            r"page \d+", r"answers to questions", r"stopped"
        ]
        if len(title_clean) < 3:
            return False
        for pat in ignore_patterns:
            if re.search(pat, title_lower):
                return False
        # Accept only if title contains some alphabetic words (not only numbers)
        if not re.search(r"[a-zA-Z]", title_clean):
            return False
        # Title should not start with a quote or look like a fragment
        if title_clean.startswith('"') or title_clean.startswith("'"):
            return False
        # Title should start with a capital letter (after stripping)
        if not title_clean or not title_clean[0].isupper():
            return False
        # Title should not be too long
        if len(title_clean) > 60:
            return False
        return True
    
    if not section_positions:
        # Fallback: treat the whole text as one section
        return [Document(page_content=text, metadata={"section": "Full Content", "type": "custom", "number": "1"})]
    
    sections = []
    seen_numbers = set()
    for i in range(len(section_positions)):
        start_pos = section_positions[i][0]
        end_pos = section_positions[i+1][0] if i+1 < len(section_positions) else len(text)
        section_content = text[start_pos:end_pos].strip()
        section_num = section_positions[i][1]
        section_title = section_positions[i][2].strip()
        sec_type = section_type if section_type != "any" else "Section"
        section_label = f"{sec_type.capitalize()} {section_num} {section_title}".strip()
        # Deduplicate by chapter number
        if section_num in seen_numbers:
            continue
        if is_relevant(section_title, section_num):
            sections.append(Document(
                page_content=section_content,
                metadata={
                    "section": section_label,
                    "type": sec_type.lower(),
                    "number": section_num
                }
            ))
            seen_numbers.add(section_num)
    return sections

def preview_sections(sections, max_preview=5):
    print("\nExtracted Sections Preview:")
    for idx, sec in enumerate(sections):
        title = sec.metadata.get("section", "Unknown")
        preview_text = sec.page_content.strip().replace("\n", " ")[:80]
        print(f"[{idx+1}] {title}: {preview_text}...")
        if idx+1 >= max_preview:
            if len(sections) > max_preview:
                print(f"...and {len(sections) - max_preview} more sections.")
            break
    print()

def load_pdf_content(pdf_path):
    """Load and preprocess PDF content"""
    try:
        print("\nLoading PDF... This may take a moment.")
        loader = PyPDFLoader(pdf_path)
        pages = loader.load()
        print(f"Loaded {len(pages)} pages.")
        return pages
    except Exception as e:
        print(f"Error loading PDF: {str(e)}")
        return None

def get_content_by_pages(pages, start_page, end_page):
    """Extract content from specific page range"""
    if start_page < 1 or end_page > len(pages) or start_page > end_page:
        return None
    
    full_text = ""
    for i in range(start_page - 1, end_page):
        content = pages[i].page_content.strip()
        full_text += content + "\n\n"
    return full_text

def get_section_detection_method():
    print("\nHow would you like to detect sections?")
    print("1. Automatic (by Chapter/Unit/Module/Section/Topic headers)")
    print("2. Manual (by page numbers)")
    print("3. By custom keyword (enter your own)")
    while True:
        choice = get_safe_input("\nEnter your choice (1-3): ")
        if choice in ["1", "2", "3"]:
            return choice
        print("Invalid choice. Please try again.")

def get_page_range(total_pages):
    while True:
        try:
            print(f"\nTotal pages in document: {total_pages}")
            start = get_safe_input("Enter start page number: ", int)
            end = get_safe_input("Enter end page number: ", int)
            
            if 1 <= start <= end <= total_pages:
                return start, end
            print("Invalid page range. Please try again.")
        except ValueError:
            print(" Please enter valid numbers.")

def get_assessment_type():
    print("\nWhat would you like to generate?")
    print("1. Question Paper (descriptive questions)")
    print("2. Quiz (multiple choice questions)")
    
    while True:
        choice = get_safe_input("\nEnter your choice (1-2): ")
        if choice in ["1", "2"]:
            return "quiz" if choice == "2" else "exam"
        print(" Invalid choice. Please try again.")

def get_pdf_file():
    while True:
        try:
            print("\nAvailable PDF files in current directory:")
            pdf_files = [f for f in os.listdir() if f.lower().endswith('.pdf')]
            for i, file in enumerate(pdf_files, 1):
                print(f"{i}. {file}")
            
            choice = get_safe_input("\n📚 Enter file number or full path: ")
            
            # Check if user entered a number
            if choice.isdigit() and 1 <= int(choice) <= len(pdf_files):
                pdf_file = pdf_files[int(choice)-1]
            else:
                pdf_file = choice.strip('"').strip("'")
            
            # Convert to absolute path if needed
            if not os.path.isabs(pdf_file):
                pdf_file = os.path.join(os.getcwd(), pdf_file)
            
            if os.path.exists(pdf_file):
                print(f"Selected: {os.path.basename(pdf_file)}")
                return pdf_file
            print(" File not found. Please try again.")
        except Exception as e:
            print(f"Error: {str(e)}")
            print("Please try again.")

def get_quiz_pattern():
    print("\n=== Quiz Pattern ===")
    print("Enter details for your quiz:")
    
    while True:
        try:
            num_questions = get_safe_input("Total number of MCQs: ", int)
            if num_questions <= 0:
                print(" Number of questions must be positive.")
                continue
                
            marks_per_question = get_safe_input("Marks per question: ", int)
            if marks_per_question <= 0:
                print(" Marks must be positive.")
                continue
                
            return num_questions, marks_per_question
        except Exception as e:
            print(f"Error: {str(e)}")

def get_section_type():
    print("\nDetect sections by:")
    print("1. Any (Chapters/Units/Modules/Sections/Topics)")
    print("2. Chapters only")
    print("3. Units only")
    print("4. Modules only")
    print("5. Sections only")
    print("6. Topics only")
    while True:
        choice = get_safe_input("\nEnter your choice (1-6): ")
        if choice == "1":
            return "any"
        elif choice == "2":
            return "chapter"
        elif choice == "3":
            return "unit"
        elif choice == "4":
            return "module"
        elif choice == "5":
            return "section"
        elif choice == "6":
            return "topic"
        print(" Invalid choice. Please try again.")

def get_course_details():
    print("\n=== Course Details ===")
    course = get_safe_input(" Course Name: ")
    code = get_safe_input(" Course Code: ")
    exam = get_safe_input(" Exam Title (e.g., Mid Sem): ")
    return course, code, exam

def get_question_pattern():
    print("\n=== Question Pattern ===")
    print("Enter number of questions and marks for each cognitive level:")
    
    levels = {
        "Remember": (" Remember (Knowledge)", "recall, define, state"),
        "Understand": (" Understand (Comprehension)", "explain, describe, discuss"),
        "Apply": (" Apply (Application)", "solve, implement, demonstrate")
    }
    
    pattern = {}
    for level, (display, examples) in levels.items():
        print(f"\n{display}")
        print(f"Examples: {examples}")
        count = get_safe_input(f"Number of questions: ", int)
        marks = get_safe_input(f"Marks per question: ", int)
        pattern[level] = (count, marks)
    
    return pattern

def get_safe_input(prompt, input_type=str, default=None):
    """Safely get user input with type conversion"""
    while True:
        try:
            value = input(prompt).strip()
            if not value and default is not None:
                return default
            if input_type == int:
                return int(value)
            return value
        except ValueError:
            print(" Please enter a valid value.")
        except Exception as e:
            print(f" Error: {str(e)}")

def handle_groq_error(func):
    """Decorator to handle Groq API errors"""
    def wrapper(*args, **kwargs):
        max_retries = 3
        for attempt in range(max_retries):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                if "rate_limit_exceeded" in str(e):
                    wait_time = (attempt + 1) * 5
                    print(f"\n Rate limit exceeded. Waiting {wait_time} seconds...")
                    time.sleep(wait_time)
                    continue
                raise e
        print(" Maximum retries exceeded. Please try again later.")
        return None
    return wrapper

def generate_questions(content: str, level: str, count: int, marks: int):
    """
    Generate unique and well-formatted questions based on cognitive level using Groq LLM.
    Args:
        content (str): The content to generate questions from
        level (str): Cognitive level ('Remember', 'Understand', or 'Apply')
        count (int): Number of questions to generate
        marks (int): Marks per question
    Returns:
        List[Dict[str, str]]: List of dictionaries containing questions and their levels
    """
    prompt = f"""
Generate {count} unique, clear, and exam-appropriate questions at the '{level}' cognitive level, each worth {marks} marks.
- Focus only on the subject matter, not on chapters, page numbers, or book structure.
- Each question must be standalone and suitable for a typical university exam paper.
- Do NOT include any instructions, meta-descriptions, or references to chapters, pages, topics, or 'here are' statements.
- Only output the questions themselves, as a numbered list, one per line.
Content:
{content}
"""
    forbidden_phrases = [
        "here are", "page number", "chapter", "section", "discussed in", "as discussed", "as described", "as explained", "as outlined", "as shown", "as mentioned", "refer to", "see section", "book", "textbook"
    ]
    try:
        response = llm.invoke(prompt)
        output = response.content if hasattr(response, 'content') else str(response)
        # Split and filter questions
        questions = [q.strip() for q in output.split('\n') if q.strip()]
        clean_questions = []
        for q in questions:
            q_lower = q.lower()
            if any(phrase in q_lower for phrase in forbidden_phrases):
                continue
            clean_questions.append(q)
            if len(clean_questions) == count:
                break
        # If not enough clean questions, fill with stubs
        while len(clean_questions) < count:
            clean_questions.append(f"[Placeholder] Exam-appropriate question for {level} level.")
        return [{"text": q, "level": level} for q in clean_questions[:count]]
    except Exception as e:
        raise Exception(f"Error generating questions: {str(e)}")

def generate_mcq(content, num_questions):
    """
    Generate MCQs from content using the LLM (ChatGroq), suitable for university-level exams.
    Returns a list of MCQ strings in the required format.
    """
    from langchain_groq import ChatGroq
    prompt = f"""
Generate {num_questions} unique, exam-appropriate multiple choice questions (MCQs) from the following content. Each question must be clear, factually correct, and suitable for a university examination—avoid trivial, vague, or overly simple questions. Do NOT reference chapters, sections, or meta-descriptions. Ensure the questions focus on core concepts, application, and analysis relevant to the subject matter. Each question should have 4 options (A, B, C, D) and indicate the correct answer. Only output the questions in the following format:
Q) <question>
A) <option>
B) <option>
C) <option>
D) <option>
CORRECT: <A/B/C/D>
Content:
"""
    prompt += content
    llm = ChatGroq(temperature=0.3, model_name="llama3-70b-8192")
    response = llm.invoke(prompt)
    # FIX: get the string content from the response object
    if hasattr(response, 'content'):
        output = response.content
    else:
        output = str(response)
    mcqs_raw = output.split('Q)')
    mcqs = []
    for mcq in mcqs_raw:
        mcq = mcq.strip()
        if not mcq:
            continue
        mcq_full = 'Q) ' + mcq
        # Only add if it contains all options and CORRECT
        if all(x in mcq_full for x in ['A)', 'B)', 'C)', 'D)', 'CORRECT:']):
            mcqs.append(mcq_full)
    # If less than requested, fill with stubs
    while len(mcqs) < num_questions:
        mcqs.append("Q) Example MCQ suitable for university exam.\nA) Option 1\nB) Option 2\nC) Option 3\nD) Option 4\nCORRECT: A")
    return mcqs[:num_questions]

def generate_exam_docx(course_name, course_code, exam_title, questions, pattern):
    doc = DocxDocument()

    # Header Section
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run("THIAGARAJAR COLLEGE OF ENGINEERING, MADURAI 625 015.\n").bold = True
    header.add_run("Department of Computer Science & Business Systems\n").bold = True
    
    # Test Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"{exam_title}\n").bold = True
    
    # Course Details Table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Fill course details
    details = [
        ("Course Code", course_code),
        ("Course Name", course_name),
        ("Degree", "B.Tech"),
        ("Programme", "CSBS"),
        ("Date", datetime.now().strftime("%d-%m-%Y (AN)"))
    ]
    
    for i, (key, value) in enumerate(details):
        cell = table.cell(i, 0)
        cell.text = key
        cell.paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value
    
    # Add Duration and Max Marks
    duration_table = doc.add_table(rows=1, cols=4)
    duration_table.style = 'Table Grid'
    cells = duration_table.rows[0].cells
    cells[0].text = "Duration"
    cells[1].text = "1 hour 45 minutes"
    cells[2].text = "Max. Marks"
    cells[3].text = str(sum(count * marks for _, (count, marks) in pattern.items()))
    
    # Answer All Questions header
    doc.add_paragraph().add_run("\nAnswer All Questions").bold = True
    doc.add_paragraph("")

    # Custom Section Headers and Improved Spacing
    section_names = {
        "Remember": "Section A: Knowledge-Based Questions",
        "Understand": "Section B: Comprehension & Understanding",
        "Apply": "Section C: Application & Analysis"
    }
    # Flatten all questions in order of pattern for continuous numbering
    ordered_questions = []
    for level, (count, marks) in pattern.items():
        level_questions = [q for q in questions if q["level"] == level][:count]
        ordered_questions.append((level, count, marks, level_questions))
    question_num = 1
    for idx, (level, count, marks, level_questions) in enumerate(ordered_questions):
        # Add section header
        section_header = doc.add_paragraph()
        section_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        section_header.add_run(section_names.get(level, f"Section {chr(65+idx)}")).bold = True
        section_header.add_run(f"   [{count} x {marks} = {count * marks} marks]").italic = True
        doc.add_paragraph("")
        # Add questions with continuous numbering
        for q in level_questions:
            question = doc.add_paragraph()
            question.paragraph_format.space_after = 12
            question.add_run(f"{question_num}. ").bold = True
            question.add_run(f"{q['text']} ")
            question.add_run(f"({marks} marks)")
            question_num += 1
        doc.add_paragraph("")  # Extra spacing after section
    # Save the document
    filename = "Generated_Question_Paper.docx"
    doc.save(filename)
    print(f"Exam generated: {filename}")
    return filename

# === Main Program ===
if __name__ == "__main__":
    print_banner()

    # Load PDF
    while True:
        pdf_file = get_pdf_file()
        pages = load_pdf_content(pdf_file)
        if pages:
            break
        print("Please try a different PDF file.")

    # Choose assessment type
    assessment_type = get_assessment_type()

    # Get content based on detection method
    detection_method = get_section_detection_method()
    if detection_method == "2":
        # Manual page selection
        start_page, end_page = get_page_range(len(pages))
        full_text = get_content_by_pages(pages, start_page, end_page)
        documents = [Document(page_content=full_text, metadata={"section": f"Pages {start_page}-{end_page}", "type": "custom"})]
    elif detection_method == "3":
        # Custom keyword
        custom_keyword = get_safe_input("Enter the custom keyword to split sections: ")
        full_text = "\n\n".join([page.page_content for page in pages])
        documents = extract_sections(full_text, section_type="custom", custom_keyword=custom_keyword)
    else:
        # Automatic section detection
        full_text = "\n\n".join([page.page_content for page in pages])
        section_type = get_section_type()
        print(f"\nAnalyzing document structure...")
        documents = extract_sections(full_text, section_type)

    print(f" Content prepared successfully!")

    # Preview sections
    preview_sections(documents)

    # Show available sections
    print("\n=== Available Sections ===")
    for i, doc in enumerate(documents):
        header = doc.page_content.split('\n')[0].strip()
        print(f"{i+1}. {doc.metadata['section']}: {header[:100]}...")

    # Get section selection
    while True:
        try:
            selected = get_safe_input("\nEnter section numbers (comma-separated): ")
            selected_indices = [int(i.strip()) - 1 for i in selected.split(",")]
            if all(0 <= i < len(documents) for i in selected_indices):
                selected_docs = [documents[i] for i in selected_indices]
                break
            raise ValueError()
        except:
            print(" Invalid selection. Please try again.")

    # Get course details
    course, code, exam = get_course_details()

    # Setup LLM and retriever
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    vectorstore = FAISS.from_documents(selected_docs, embeddings)
    retriever = vectorstore.as_retriever()
    qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)

    if assessment_type == "quiz":
        # Generate quiz as a question paper with MCQs only
        num_questions, marks_per_question = get_quiz_pattern()
        docx = DocxDocument()
        docx.add_heading(f"{exam} Quiz", 0)
        docx.add_paragraph(f"Course Name: {course}")
        docx.add_paragraph(f"Course Code: {code}")
        docx.add_paragraph(f"Total Marks: {num_questions * marks_per_question}")
        docx.add_paragraph("")
        question_num = 1
        for doc in selected_docs:
            docx.add_heading(doc.metadata['section'], level=1)
            questions_per_section = num_questions // len(selected_docs)
            mcqs = generate_mcq(doc.page_content, questions_per_section)
            for mcq in mcqs:
                lines = mcq.split('\n')
                question = lines[0].lstrip('Q) ')
                options = lines[1:5]
                correct = lines[5].lstrip('CORRECT: ') if len(lines) > 5 else 'A'
                q_par = docx.add_paragraph(f"{question_num}. {question}")
                for opt in options:
                    docx.add_paragraph(opt.strip(), style='List Bullet')
                docx.add_paragraph(f"[{marks_per_question} marks] (Recommended time: 60 seconds)")
                docx.add_paragraph("")
                question_num += 1
        # Add answer key at the end
        docx.add_page_break()
        docx.add_heading("Answer Key", level=1)
        docx.add_paragraph("(For instructor use only)")
        question_num = 1
        for doc in selected_docs:
            mcqs = generate_mcq(doc.page_content, num_questions // len(selected_docs))
            for mcq in mcqs:
                lines = mcq.split('\n')
                if len(lines) > 5:
                    correct = lines[5].lstrip('CORRECT: ')
                    docx.add_paragraph(f"Q{question_num}: {correct}")
                question_num += 1
        output_file = "Generated_Quiz.docx"
        docx.save(output_file)
        print(f"\nQuiz generated: {output_file}")
    else:
        # Generate regular question paper
        pattern = get_question_pattern()
        
        questions = []
        total_marks = 0
        for doc in selected_docs:
            for level, (count, marks) in pattern.items():
                qs = generate_questions(doc.page_content, level, count, marks)
                questions.extend(qs)
                total_marks += count * marks
        
        generate_exam_docx(course, code, exam, questions, pattern)
        output_file = "Generated_Question_Paper.docx"
        print(f"\nQuestion paper generated: {output_file}")
